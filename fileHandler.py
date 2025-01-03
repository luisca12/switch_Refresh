from log import authLog
from docx import Document
from docx.shared import RGBColor, Pt
from auth import Auth
from commandsCLI import shCoreInfo
import openpyxl

import re
import os
import csv
import json
import traceback
import ipaddress

def docxWorkstationsCore(validIPs, username, netDevice):
    commandOutput = shCoreInfo(validIPs, username, netDevice)

    try: 
        wordFile = "Template 9300 Layer 3 Core with Workstations, Phones or Printers Config.docx"
        wordDOC = Document(wordFile)
        authLog.info(f"file successfully found: {wordFile}")
        print(f"INFO: file successfully found: {wordFile}.")

        


    except Exception as error:
        print(f"ERROR: {error}\n", traceback.format_exc())
        authLog.error(f"Wasn't possible to choose the DOCX file, error message: {error}\n{traceback.format_exc()}")

def docxCore(validIPs, username, netDevice):
    commandOutput = shCoreInfo(validIPs, username, netDevice)

    try: 
        wordFile = "Template 9300 Layer 3 Core without Workstations, Phones or Printers Config.docx"
        wordDOC = Document(wordFile)
        authLog.info(f"file successfully found: {wordFile}")
        print(f"INFO: file successfully found: {wordFile}.")

    except Exception as error:
        print(f"ERROR: {error}\n", traceback.format_exc())
        authLog.error(f"Wasn't possible to choose the DOCX file, error message: {error}\n{traceback.format_exc()}")

def docxIDF(validDeviceIPs, username, netDevice):

    try: 
        for validIPs in validDeviceIPs:
            commandOutput = []
            commandOutput, oldInt, oldIntDesc, vlanIDList, oldIntStat = shCoreInfo(validIPs, username, netDevice)
            wordFile = "Template 9300 Layer 2 IDF Config.docx"
            wordDOC = Document(wordFile)
            authLog.info(f"file successfully found: {wordFile}")
            print(f"INFO: file successfully found: {wordFile}.")

            for command in commandOutput:
                docPara = wordDOC.add_paragraph()
                paraRun = docPara.add_run(command)
                paraRun.bold = True
                authLog.info(f"The string {command} was appended to the file {wordFile} for device {validIPs}")
            
            validIPsStr = ''.join(validIPs)
            wordDOC.save(f'Outputs/IDF Refresh for {validIPsStr}.docx')
        
            cutSheet(validIPs, oldInt, oldIntDesc, vlanIDList, oldIntStat)

    except Exception as error:
        print(f"ERROR: {error}\n", traceback.format_exc())
        authLog.error(f"Wasn't possible to choose the DOCX file, error message: {error}\n{traceback.format_exc()}")

def cutSheet(validIPs, *args):
    cutSheetFile = "cutSheet Base.xlsx"
    outputFolder = "Outputs"

    CutSheetBase = openpyxl.load_workbook(cutSheetFile)
    authLog.info(f"Automation successfully found the XLSX: {cutSheetFile}")
    print(f"INFO: Successfully found the XLSX: {cutSheetFile}")
    CutSheetSheet = CutSheetBase.active

    newCutSheetFile = f'{validIPs} CutSheet.xlsx'
    CutSheetSheet['A1'] = f'{validIPs} CutSheet'
    authLog.info(f"Automation successfully generated the name for the new cut sheet: {newCutSheetFile}")

    columns = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J']
    authLog.info(f"These are the columns that will be modified: {columns}")
    
    # Iterate over the arguments (lists) and the columns (A, B, C, D)
    for i, dataList in enumerate(args):
        column = columns[i]  # Get the corresponding column (A, B, C, D, etc.)
        authLog.info(f"The script is currently on column: {column}")
        for row, value in enumerate(dataList, start=3):  # Start from row 3
            authLog.info(f"The script is currently on row: {row}")
            cell = f'{column}{row}'  # Create the cell reference (e.g., A3, B3, etc.)
            CutSheetSheet[cell] = value  # Paste the data in the cell
            authLog.info(f"The new value on cell: {cell} is: {value}")
            print(f"INFO: The new value on cell: {cell} is: {value}")

    newCutSheet = os.path.join(outputFolder, newCutSheetFile)
    authLog.info(f"New cut sheet is: {newCutSheet}")
    CutSheetBase.save(newCutSheet)
    authLog.info(f"Cut sheet was successfully saved")
    print(f"INFO: New cut sheet is: {newCutSheet}")
    print(f"INFO: Cut sheet was successfully saved")